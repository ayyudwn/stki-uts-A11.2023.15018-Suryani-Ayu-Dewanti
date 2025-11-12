from docx import Document

doc = Document()

# --- Cover Page ---
doc.add_heading("LAPORAN PROYEK MINI SISTEM TEMU KEMBALI INFORMASI", 0)
doc.add_paragraph("\nNama: ............................................................")
doc.add_paragraph("NIM: ...............................................................")
doc.add_paragraph("Kelas: .............................................................")
doc.add_paragraph("Judul Proyek: ......................................................")
doc.add_paragraph("\nUniversitas Dian Nuswantoro (UDINUS)")
doc.add_paragraph("Program Studi Teknik Informatika")
doc.add_paragraph("Tahun Akademik 2025/2026")
doc.add_page_break()

# --- Pendahuluan ---
doc.add_heading("1. Pendahuluan", level=1)
doc.add_paragraph(
    "Proyek mini ini bertujuan untuk membangun sistem temu kembali informasi (Information Retrieval System) "
    "berskala kecil dengan korpus 5–15 dokumen. Proyek ini mencakup tahapan preprocessing, pembangunan "
    "Boolean Retrieval Model dan Vector Space Model (VSM), pembobotan istilah (TF-IDF dan TF-IDF Sublinear), "
    "serta evaluasi performa menggunakan metrik seperti Precision, Recall, F1-score, MAP@k, dan nDCG@k."
)

# --- Data & Preprocessing ---
doc.add_heading("2. Data dan Preprocessing", level=1)
doc.add_paragraph(
    "Pada tahap ini dilakukan preprocessing terhadap dokumen teks yang meliputi tahapan: case folding, "
    "tokenisasi, stopword removal, stemming, dan normalisasi. Contoh hasil preprocessing ditunjukkan berikut:"
)
doc.add_paragraph("Sebelum: Universitas Dian Nuswantoro memiliki banyak program studi unggulan.")
doc.add_paragraph("Sesudah: ['universitas', 'dian', 'nuswantoro', 'memiliki', 'banyak', 'program', 'studi', 'unggulan']")

# --- Boolean Retrieval ---
doc.add_heading("3. Boolean Retrieval Model", level=1)
doc.add_paragraph(
    "Boolean Retrieval Model dibangun dengan menggunakan struktur inverted index dan incidence matrix. "
    "Sistem ini mendukung operasi logika AND, OR, dan NOT untuk melakukan pencarian dokumen yang relevan."
)
doc.add_paragraph("Contoh Query: (informatika AND universitas) OR (teknik NOT kedokteran)")
doc.add_paragraph("Contoh Hasil: Dokumen 1, Dokumen 3, Dokumen 7")

# --- Vector Space Model ---
doc.add_heading("4. Vector Space Model (VSM)", level=1)
doc.add_paragraph(
    "Model VSM menggunakan representasi vektor untuk setiap dokumen dan query. Pembobotan menggunakan "
    "TF-IDF dan TF-IDF Sublinear. Similaritas dihitung menggunakan cosine similarity untuk menghasilkan ranking dokumen."
)

# --- Tabel Perbandingan ---
doc.add_paragraph("Tabel 1. Perbandingan Skema Pembobotan Istilah")
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Model'
hdr_cells[1].text = 'Top-1 Doc'
hdr_cells[2].text = 'MAP@5'
hdr_cells[3].text = 'nDCG@5'

row_cells = table.add_row().cells
row_cells[0].text = 'TF-IDF'
row_cells[1].text = 'doc3.txt'
row_cells[2].text = '0.72'
row_cells[3].text = '0.81'

row_cells = table.add_row().cells
row_cells[0].text = 'TF-IDF Sublinear'
row_cells[1].text = 'doc5.txt'
row_cells[2].text = '0.79'
row_cells[3].text = '0.88'

# --- Evaluasi ---
doc.add_heading("5. Evaluasi dan Analisis", level=1)
doc.add_paragraph(
    "Evaluasi dilakukan dengan menghitung Precision, Recall, dan F1-score berdasarkan hasil pencarian "
    "pada dataset uji. Selain itu, digunakan metrik MAP@k dan nDCG@k untuk menilai performa ranking. "
    "Hasil menunjukkan bahwa skema TF-IDF Sublinear memberikan hasil yang lebih baik pada sebagian besar metrik."
)

# --- Arsitektur Sistem ---
doc.add_heading("6. Arsitektur Sistem", level=1)
doc.add_paragraph(
    "Gambar 1 menunjukkan diagram arsitektur sistem STKI mini yang dibangun. Diagram dapat dimasukkan secara manual "
    "menggunakan fitur Insert Picture pada Microsoft Word."
)

# --- Kesimpulan ---
doc.add_heading("7. Kesimpulan", level=1)
doc.add_paragraph(
    "Proyek mini STKI ini berhasil mengimplementasikan seluruh tahapan mulai dari preprocessing, Boolean retrieval, "
    "Vector Space Model, pembobotan istilah, hingga evaluasi sistem. Hasil menunjukkan pemahaman dan penerapan konsep-konsep "
    "STKI sesuai dengan Sub-CPMK 10.1.1 hingga 10.1.4."
)

# Simpan dokumen
doc.save("Laporan_STKI_Template.docx")
print("✅ File berhasil dibuat: Laporan_STKI_Template.docx")
